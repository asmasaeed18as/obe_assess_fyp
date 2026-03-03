import io
import re
import fitz  # PyMuPDF
from docx import Document

try:
    import pymupdf4llm  # Optional dependency for Markdown extraction
except Exception:
    pymupdf4llm = None


def _extract_clo_relevant_text(text):
    """
    Keep only CLO-relevant sections from noisy extracted text.
    Falls back to original text if no relevant sections are found.
    """
    if not text:
        return ""

    lines = [line.strip() for line in text.splitlines()]
    keyword_pattern = re.compile(
        r"\b("
        r"clo(?:[-\s]?\d+)?|"
        r"course learning outcome(?:s)?|"
        r"plo(?:[-\s]?\d+)?|"
        r"programme learning outcome(?:s)?|"
        r"program learning outcome(?:s)?|"
        r"bt(?:\s*level)?|"
        r"bloom(?:'s)?|"
        r"cognitive(?:\s*domain)?|"
        r")\b",
        re.IGNORECASE,
    )

    selected_lines = []
    seen = set()
    line_count = len(lines)

    for idx, line in enumerate(lines):
        if not line:
            continue
        if keyword_pattern.search(line):
            start = max(0, idx - 1)
            end = min(line_count, idx + 2)
            for ctx_idx in range(start, end):
                ctx_line = lines[ctx_idx].strip()
                if not ctx_line:
                    continue
                if ctx_line not in seen:
                    seen.add(ctx_line)
                    selected_lines.append(ctx_line)

    if selected_lines:
        return "\n".join(selected_lines)

    return text


def extract_text_from_file(file_obj, filename, only_clo_relevant=True):
    """
    State-of-the-Art LLM Extraction:
    Converts PDFs and DOCX files into pure Markdown.
    LLMs natively understand Markdown tables, completely eliminating hallucination.
    """
    file_ext = filename.lower().split('.')[-1]

    try:
        file_obj.seek(0)
        file_bytes = file_obj.read()

        # ==========================================
        # 1. PDF EXTRACTION (The LLM Standard)
        # ==========================================
        if file_ext == 'pdf':
            # Open the PDF stream directly from memory
            doc = fitz.open(stream=file_bytes, filetype="pdf")

            if pymupdf4llm is not None:
                md_text = pymupdf4llm.to_markdown(doc)
            else:
                extracted = []
                for page in doc:
                    extracted.append(page.get_text("text"))
                md_text = "\n".join(extracted)

            if only_clo_relevant:
                return _extract_clo_relevant_text(md_text)
            return md_text

        # ==========================================
        # 2. DOCX EXTRACTION (Formatted as Markdown)
        # ==========================================
        elif file_ext in ['docx', 'doc']:
            doc = Document(io.BytesIO(file_bytes))
            extracted_text = []

            # Extract paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    extracted_text.append(p.text.strip())

            # Extract tables and format them as Markdown tables so the LLM understands them
            for table in doc.tables:
                extracted_text.append("\n")  # Add spacing before table
                for i, row in enumerate(table.rows):
                    clean_row = []
                    for cell in row.cells:
                        cell_text = cell.text.replace('\n', ' ').strip()
                        if not cell_text:
                            continue
                        # Keep markdown parser stable when source cell contains pipe characters.
                        clean_row.append(cell_text.replace("|", "/"))
                    if clean_row:
                        # Format as a Markdown row: | Column 1 | Column 2 |
                        extracted_text.append("| " + " | ".join(clean_row) + " |")

                        # Add the Markdown header separator after the first row
                        if i == 0:
                            extracted_text.append("|" + "|".join(["---"] * len(clean_row)) + "|")
                extracted_text.append("\n")

            final_text = "\n".join(extracted_text)
            if only_clo_relevant:
                return _extract_clo_relevant_text(final_text)
            return final_text

        else:
            raise ValueError("Unsupported file format.")

    except Exception as e:
        print(f"Extraction Error: {str(e)}")
        return ""
