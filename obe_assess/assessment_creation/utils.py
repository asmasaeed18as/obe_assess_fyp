# assessment_creation/utils.py
import fitz   # PyMuPDF
import io
import zipfile

# ReportLab Imports (PDF Generation)
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch

# Python-Docx Imports (Word Generation)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def extract_text_from_pdf_filefield(django_file) -> str:
    """
    Extracts text from a Django UploadedFile/FileField using PyMuPDF (fitz).
    """
    try:
        file_bytes = django_file.read()
        text = ""
        # Open with fitz using stream
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""

# ==========================================
# 1. DOCX GENERATOR
# ==========================================
def generate_docx_assessment(questions, title, content_type):
    """
    Generates a Word Document based on content_type filters.
    content_type options: 'full', 'questions_only', 'rubrics_only', 'answers_only'
    """
    doc = Document()
    
    # --- Document Title ---
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Subtitle with Format Type
    doc.add_paragraph(f"Format: {content_type.replace('_', ' ').title()}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("_" * 50).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for idx, q in enumerate(questions):
        # --- Question Header (Always visible) ---
        p = doc.add_paragraph()
        p.space_before = Pt(12)
        runner = p.add_run(f"Q{idx + 1}: {q.get('question', 'Question text missing')}")
        runner.bold = True
        runner.font.size = Pt(11)

        # --- Filter 1: Metadata (Visible in Full Doc & Question Paper) ---
        if content_type in ['full', 'questions_only']:
            meta = q.get('meta', {})
            # Safely get values with defaults
            marks = q.get('marks', '0')
            clo = meta.get('clo', '-')
            bloom = meta.get('bloom', '-')
            
            meta_text = f"Marks: {marks} | CLO: {clo} | Bloom: {bloom}"
            meta_p = doc.add_paragraph(meta_text)
            meta_p.runs[0].italic = True
            meta_p.runs[0].font.color.rgb = RGBColor(100, 100, 100) # Grey color

        # --- Filter 2: Model Answers (Visible in Full Doc & Answer Key) ---
        if content_type in ['full', 'answers_only']:
            doc.add_paragraph("Model Answer:", style='Heading 3')
            doc.add_paragraph(q.get('answer', 'N/A'))

        # --- Filter 3: Rubrics (Visible in Full Doc & Rubric Sheet) ---
        if content_type in ['full', 'rubrics_only']:
            rubric = q.get('rubric', {})
            if rubric and isinstance(rubric, dict):
                doc.add_paragraph("Grading Rubric:", style='Heading 3')
                
                # Create Table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Header Row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Criteria'
                hdr_cells[1].text = 'Description'
                
                # Make header bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Data Rows
                for key, value in rubric.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)
            
            doc.add_paragraph("") # Spacer

    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 2. PDF GENERATOR
# ==========================================
def generate_pdf_assessment(questions, title, content_type):
    """
    Generates a PDF Document based on content_type filters using ReportLab.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # --- Title Section ---
    story.append(Paragraph(title, styles['Title']))
    story.append(Paragraph(f"Format: {content_type.replace('_', ' ').title()}", styles['Normal']))
    story.append(Spacer(1, 20))

    # --- Custom Styles ---
    style_q = ParagraphStyle('Question', parent=styles['Heading3'], spaceAfter=5, textColor=colors.darkblue)
    style_meta = ParagraphStyle('Meta', parent=styles['Italic'], fontSize=9, textColor=colors.gray, leftIndent=10)
    style_ans = ParagraphStyle('Answer', parent=styles['Normal'], backColor=colors.whitesmoke, borderPadding=5, leftIndent=10)

    for idx, q in enumerate(questions):
        # 1. Question Text
        q_text = f"Q{idx + 1}: {q.get('question', 'Question text missing')}"
        story.append(Paragraph(q_text, style_q))

        # 2. Metadata (Marks, CLO)
        if content_type in ['full', 'questions_only']:
            meta = q.get('meta', {})
            marks = q.get('marks', '0')
            clo = meta.get('clo', '-')
            bloom = meta.get('bloom', '-')
            
            meta_info = f"Marks: {marks} | CLO: {clo} | Bloom: {bloom}"
            story.append(Paragraph(meta_info, style_meta))
        
        story.append(Spacer(1, 6))

        # 3. Model Answer
        if content_type in ['full', 'answers_only']:
            story.append(Paragraph("<b>Model Answer:</b>", styles['Normal']))
            story.append(Paragraph(q.get('answer', 'N/A'), style_ans))
            story.append(Spacer(1, 10))

        # 4. Rubrics
        if content_type in ['full', 'rubrics_only']:
            rubric = q.get('rubric', {})
            if rubric and isinstance(rubric, dict):
                story.append(Paragraph("<b>Rubric:</b>", styles['Normal']))
                story.append(Spacer(1, 4))

                # Table Data (Wrapped in Paragraphs for text wrapping)
                data = [['Criteria', 'Description']]
                for k, v in rubric.items():
                    data.append([Paragraph(str(k), styles['Normal']), Paragraph(str(v), styles['Normal'])])
                
                # Table Styling
                col_widths = [1.5 * inch, 4.5 * inch]
                t = Table(data, colWidths=col_widths)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BOX', (0, 0), (-1, -1), 1, colors.black),
                    ('PADDING', (0, 0), (-1, -1), 6),
                ]))
                story.append(t)
        
        story.append(Spacer(1, 15))
        story.append(Paragraph("_" * 60, styles['Normal']))
        story.append(Spacer(1, 15))

    doc.build(story)
    buffer.seek(0)
    return buffer

# ==========================================
# 3. ZIP GENERATOR (THE BUNDLER)
# ==========================================
def generate_zip_bundle(assessment, file_format='docx'):
    """
    Creates a ZIP file containing 3 separate documents:
    1. Question Paper (No answers)
    2. Model Answers (No rubrics/meta)
    3. Grading Rubrics (No answers)
    """
    zip_buffer = io.BytesIO()
    
    # Get questions from the DB object
    questions = assessment.result_json.get('questions', [])
    base_title = f"{assessment.assessment_type} - {assessment.material.title}"

    # Define the 3 separate files we want to create
    docs_to_generate = [
        ("Question_Paper", "questions_only"),
        ("Model_Answers", "answers_only"),
        ("Grading_Rubrics", "rubrics_only")
    ]

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for filename_suffix, content_type in docs_to_generate:
            file_name = f"{filename_suffix}.{file_format}"
            
            # Generate specific content
            if file_format == 'docx':
                file_buffer = generate_docx_assessment(questions, base_title, content_type)
            else:
                # Default to PDF if not docx
                file_buffer = generate_pdf_assessment(questions, base_title, content_type)
            
            # Write file to zip
            zip_file.writestr(file_name, file_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer